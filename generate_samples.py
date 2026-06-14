"""
generate_samples.py
-------------------
Utility script to generate sample test files in the sample_files/ directory.
Run this once to populate the directory before testing the analyzer.

Usage:
    python generate_samples.py
"""

import json
import os
from datetime import datetime
from pathlib import Path


SAMPLE_DIR = Path(__file__).parent / "sample_files"


def create_sample_docx() -> None:
    """Generate a sample .docx with realistic core properties."""
    try:
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
    except ImportError:
        print("[!] python-docx not installed. Skipping DOCX sample.")
        return

    doc = Document()

    # Core properties
    props = doc.core_properties
    props.title = "Forensic Investigation Report"
    props.subject = "Digital Evidence Analysis"
    props.author = "Jane Investigator"
    props.keywords = "forensics, DFIR, metadata, evidence"
    props.description = "Sample document for metadata analyzer testing."

    doc.add_heading("Forensic Investigation Report", level=0)
    doc.add_paragraph(
        "This is a sample Word document generated for testing the Metadata Analyzer tool. "
        "It contains realistic core properties including author, timestamps, and revision data."
    )
    doc.add_heading("Evidence Summary", level=1)
    doc.add_paragraph("Exhibit A: Hard drive image (SHA-256 hash verified)")
    doc.add_paragraph("Exhibit B: Network packet capture (PCAP format)")
    doc.add_paragraph("Exhibit C: System logs from compromised host")

    out = SAMPLE_DIR / "sample_report.docx"
    doc.save(str(out))
    print(f"[+] Created: {out}")


def create_sample_pdf() -> None:
    """Generate a minimal valid PDF with metadata using raw PDF syntax."""
    # Build a minimal PDF by hand — no dependency required
    now = datetime.now().strftime("D:%Y%m%d%H%M%S")
    content = f"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj

2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj

3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
   /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj

4 0 obj
<< /Length 44 >>
stream
BT /F1 12 Tf 72 720 Td (Metadata Analyzer Sample) Tj ET
endstream
endobj

5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj

6 0 obj
<< /Title (Metadata Analyzer Test Document)
   /Author (Michael O.)
   /Creator (Metadata Analyzer Sample Generator)
   /Producer (Python raw PDF writer)
   /Subject (Digital Forensics Portfolio)
   /Keywords (forensics metadata DFIR)
   /CreationDate ({now})
   /ModDate ({now})
>>
endobj

xref
0 7
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000266 00000 n
0000000360 00000 n
0000000441 00000 n

trailer
<< /Size 7 /Root 1 0 R /Info 6 0 R >>
startxref
700
%%EOF
"""
    out = SAMPLE_DIR / "sample_document.pdf"
    out.write_text(content, encoding="latin-1")
    print(f"[+] Created: {out}")


def create_sample_png() -> None:
    """Generate a minimal PNG file (no Pillow required — raw bytes)."""
    import struct
    import zlib

    def make_chunk(chunk_type: bytes, data: bytes) -> bytes:
        c = chunk_type + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)

    # 1x1 red pixel
    width, height = 1, 1
    raw = b"\x00\xFF\x00\x00"  # filter byte + RGBA
    compressed = zlib.compress(raw)

    png = (
        b"\x89PNG\r\n\x1a\n"
        + make_chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
        + make_chunk(b"IDAT", compressed)
        + make_chunk(b"IEND", b"")
    )

    out = SAMPLE_DIR / "sample_image.png"
    out.write_bytes(png)
    print(f"[+] Created: {out}")


def main() -> None:
    SAMPLE_DIR.mkdir(parents=True, exist_ok=True)
    print(f"[*] Generating sample files in: {SAMPLE_DIR}\n")
    create_sample_docx()
    create_sample_pdf()
    create_sample_png()
    print("\n[✓] Sample files ready. Run: python analyzer.py sample_files/")


if __name__ == "__main__":
    main()
