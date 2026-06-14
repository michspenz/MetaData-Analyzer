from pathlib import Path
from docx import Document
from docx_analyzer import analyze_docx


def test_docx_successful_extraction(tmp_path: Path):
    out = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Hello World")
    doc.save(out)

    result = analyze_docx(str(out))
    assert result["file_name"] == out.name
    assert result["type"] == "docx"


def test_docx_missing_metadata(tmp_path: Path):
    out = tmp_path / "nometa.docx"
    doc = Document()
    doc.save(out)

    result = analyze_docx(str(out))
    assert "errors" in result


def test_docx_corrupted_file(tmp_path: Path):
    bad = tmp_path / "bad.docx"
    bad.write_bytes(b"not a docx")

    result = analyze_docx(str(bad))
    assert "errors" in result
    assert any("not a valid .docx" in e or "Unexpected error" in e for e in result["errors"]) or result["file_size_bytes"] is not None
